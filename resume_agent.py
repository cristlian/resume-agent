#!/usr/bin/env python3
"""
Resume Tailoring Agent - Optimizes resumes for target job descriptions using Gemini.

Usage:
    python resume_agent.py [--jd-pdf JD.pdf] [--resume-pdf Resume.pdf] [--api-key KEY]

By default:
    - JD: fetches the latest PDF from ./JD folder (by creation time)
    - Resume: uses ./Yifei-Lian-Resume-merged.pdf
"""

import argparse
import json
import os
import re
import sys
import time
import threading
from urllib.parse import urlparse
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()
from urllib.request import Request, urlopen
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager

from google import genai
from google.genai import types
import pypdf
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

from resume_core import (
    build_critic_prompt,
    check_for_markdown as core_check_for_markdown,
    enforce_plain_text as core_enforce_plain_text,
    extract_length_constraint as core_extract_length_constraint,
    extract_readable_text_from_html,
    extract_role_from_jd_filename as core_extract_role_from_jd_filename,
    format_profile_facts_for_prompt,
    has_blocking_critic_issues,
    load_profile_facts,
    parse_structured_critic,
)

# === CONFIGURATION ===
MODEL_NAME = "gemini-3-pro-preview"
WORKSPACE_DIR = Path(__file__).parent.resolve()
DEFAULT_RESUME = WORKSPACE_DIR / "Yifei-Lian-Resume-merged.pdf"
JD_FOLDER = WORKSPACE_DIR / "JD"
TAILORED_RESUMES_FOLDER = WORKSPACE_DIR / "tailored_resumes"
APPLICATION_RESPONSES_FOLDER = WORKSPACE_DIR / "application_question_response"
DEFAULT_PROFILE_FACTS = WORKSPACE_DIR / "profile_facts.yaml"
MAX_RETRIES = 4
RETRY_BASE_SECONDS = 2
RETRY_MAX_SECONDS = 20

# === PERFORMANCE UTILITIES ===
@contextmanager
def timed_section(name: str):
    """Context manager to time and report section duration."""
    start = time.perf_counter()
    yield
    elapsed = time.perf_counter() - start
    print(f"  [TIMING] {name}: {elapsed:.2f}s")


class Spinner:
    """Minimal CLI spinner for long-running operations."""
    UNICODE_FRAMES = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]
    ASCII_FRAMES = ["-", "\\", "|", "/"]
    
    def __init__(self, message: str = "Processing"):
        self.message = message
        self._stop_event = threading.Event()
        self._thread = None
        self.frames = self._resolve_frames()

    def _resolve_frames(self) -> list[str]:
        """Choose spinner frames compatible with current stdout encoding."""
        encoding = sys.stdout.encoding or "utf-8"
        try:
            for frame in self.UNICODE_FRAMES:
                frame.encode(encoding)
            return self.UNICODE_FRAMES
        except (LookupError, UnicodeEncodeError):
            return self.ASCII_FRAMES
    
    def _spin(self):
        idx = 0
        while not self._stop_event.is_set():
            frame = self.frames[idx % len(self.frames)]
            print(f"\r  {frame} {self.message}...", end="", flush=True)
            idx += 1
            time.sleep(0.1)
        # Clear spinner line
        print(f"\r" + " " * (len(self.message) + 10) + "\r", end="", flush=True)
    
    def __enter__(self):
        self._stop_event.clear()
        self._thread = threading.Thread(target=self._spin, daemon=True)
        self._thread.start()
        return self
    
    def __exit__(self, *args):
        self._stop_event.set()
        if self._thread:
            self._thread.join(timeout=0.5)


def parallel_extract_pdfs(jd_path: Path, resume_path: Path) -> tuple[str | types.Part, str | types.Part]:
    """Extract text from both PDFs in parallel using thread pool."""
    results = {}
    
    with ThreadPoolExecutor(max_workers=2) as executor:
        futures = {
            executor.submit(load_pdf_content, jd_path): 'jd',
            executor.submit(load_pdf_content, resume_path): 'resume'
        }
        for future in as_completed(futures):
            key = futures[future]
            results[key] = future.result()
    
    return results['jd'], results['resume']


# === SYSTEM INSTRUCTIONS ===
ARCHITECT_SYSTEM_INSTRUCTION = """<role>

You are Gemini 3, a specialized assistant for Resume Optimization and ATS Alignment.

You are precise, analytical, and persistent.

</role>



<instructions>

1. **Plan**: Analyze the job description and resume variants to determine the best alignment.

2. **Execute**: Select and rephrase the most relevant achievements, skills, and experiences from the provided resume(s) to match the target role.

3. **Validate**: Ensure all extracted information is factual, user-provided, and directly matches job requirements without exaggeration or inference.

4. **Format**: Output a clean, well-structured one-page resume with standardized sections, appropriate formatting, and professional clarity.

</instructions>



<constraints>

- Verbosity: Medium

- Tone: Formal, Technical

- OUTPUT FORMAT: Plain text only. NO markdown syntax whatsoever.
   - Use UPPERCASE for section headers
   - Use simple dashes or asterisks for bullet points, not markdown-style
   - No # symbols for headings
   - No ``` code blocks
   - No **bold** or *italic* formatting

</constraints>



<output_format>

Structure your response as follows:

1. **Executive Summary**: Brief overview of tailoring logic-why certain experiences were selected and how they match the target role.

2. **Tailored Resume**:

   - Name / Phone / Email / GitHub

   - Work Experience (chronological, most relevant content prioritized)

   - Project Experience (concise, impact-focused)

   - Education (institution, degree, dates)

   - Professional Skills (tools, languages, systems)

</output_format>"""


def get_latest_jd_pdf() -> Path | None:
    """Find the most recently created PDF in the JD folder."""
    if not JD_FOLDER.exists():
        return None
    
    pdf_files = list(JD_FOLDER.glob("*.pdf"))
    if not pdf_files:
        return None
    
    # Sort by creation time (st_ctime), most recent first
    latest = max(pdf_files, key=lambda p: p.stat().st_ctime)
    return latest


def parse_args() -> argparse.Namespace:
    """Parse and validate command-line arguments."""
    parser = argparse.ArgumentParser(
        description="Tailor a resume to a job description using Gemini AI.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
Defaults:
  --jd-url      Optional URL input; if fetch fails, paste text interactively
  --jd-text     Direct JD text input (highest priority)
  --jd-pdf      Latest PDF from {JD_FOLDER}
  --resume-pdf  {DEFAULT_RESUME}
  --profile-facts  {DEFAULT_PROFILE_FACTS} (if file exists)
""",
    )
    parser.add_argument(
        "--jd-url",
        required=False,
        help="Job description URL (HTML fetch with fallback to pasted text)",
    )
    parser.add_argument(
        "--jd-text",
        required=False,
        help="Direct job description text input",
    )
    parser.add_argument(
        "--jd-pdf",
        required=False,
        help="Path to the job description PDF (fallback when URL/text not provided)",
    )
    parser.add_argument(
        "--resume-pdf",
        required=False,
        help=f"Path to the merged resume PDF (default: {DEFAULT_RESUME.name})",
    )
    parser.add_argument(
        "--api-key",
        help="Gemini API key (overrides GEMINI_API_KEY env var)",
    )
    parser.add_argument(
        "--profile-facts",
        required=False,
        help="Path to profile facts YAML (default: profile_facts.yaml if present)",
    )
    return parser.parse_args()


def _canonicalize_url_for_match(url: str) -> tuple[str, str]:
    """Return (host, normalized_path) for robust URL matching."""
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    path = re.sub(r"/{2,}", "/", parsed.path).rstrip("/").lower()
    return host, path


def _extract_ashby_job_text_via_posting_api(job_url: str, timeout_seconds: int = 15) -> str:
    """Fetch JD text for Ashby-hosted jobs via public posting API."""
    parsed = urlparse(job_url)
    segments = [segment for segment in parsed.path.split("/") if segment]
    if len(segments) < 2:
        raise ValueError("Ashby URL format not recognized.")

    job_board_name = segments[0]
    target_slug = segments[-1].lower()
    api_url = f"https://api.ashbyhq.com/posting-api/job-board/{job_board_name}"

    request = Request(
        api_url,
        headers={
            "User-Agent": "resume-optimizer/1.0 (+https://local.cli)",
            "Accept": "application/json",
        },
    )
    with urlopen(request, timeout=timeout_seconds) as response:
        data = json.loads(response.read().decode("utf-8", errors="replace"))

    jobs = data.get("jobs")
    if not isinstance(jobs, list) or not jobs:
        raise ValueError(f"No jobs found from Ashby posting API for board '{job_board_name}'.")

    target_host, target_path = _canonicalize_url_for_match(job_url)
    matched_job: dict | None = None
    fallback_match: dict | None = None
    for job in jobs:
        if not isinstance(job, dict):
            continue
        for key in ("jobUrl", "applyUrl"):
            candidate_url = job.get(key)
            if not isinstance(candidate_url, str) or not candidate_url.strip():
                continue
            candidate_host, candidate_path = _canonicalize_url_for_match(candidate_url)
            candidate_segments = [segment for segment in candidate_path.split("/") if segment]
            if candidate_host == target_host and candidate_path == target_path:
                matched_job = job
                break
            if target_slug in candidate_segments:
                fallback_match = fallback_match or job
        if matched_job is not None:
            break
    if matched_job is None:
        matched_job = fallback_match
    if matched_job is None:
        raise ValueError("Could not match target job URL in Ashby posting API response.")

    description_plain = matched_job.get("descriptionPlain")
    if isinstance(description_plain, str):
        description_text = description_plain.strip()
    else:
        description_text = ""

    if len(description_text) < 80:
        description_html = matched_job.get("descriptionHtml")
        if isinstance(description_html, str):
            description_text = extract_readable_text_from_html(description_html).strip()

    if len(description_text) < 80:
        raise ValueError("Ashby posting API returned insufficient JD description text.")

    details: list[str] = []
    for field, label in (
        ("title", "Title"),
        ("location", "Location"),
        ("department", "Department"),
        ("team", "Team"),
        ("employmentType", "Employment Type"),
    ):
        value = matched_job.get(field)
        if isinstance(value, str) and value.strip():
            details.append(f"{label}: {value.strip()}")

    details.append("")
    details.append(description_text)
    return "\n".join(details).strip()


def fetch_jd_text_from_url(url: str, timeout_seconds: int = 15) -> str:
    """Fetch a JD page and extract readable text."""
    parsed = urlparse(url)
    ashby_error: Exception | None = None

    if parsed.netloc.lower().endswith("jobs.ashbyhq.com"):
        try:
            return _extract_ashby_job_text_via_posting_api(url, timeout_seconds=timeout_seconds)
        except Exception as e:
            ashby_error = e

    try:
        request = Request(
            url,
            headers={"User-Agent": "resume-optimizer/1.0 (+https://local.cli)"},
        )
        with urlopen(request, timeout=timeout_seconds) as response:
            html_bytes = response.read()

        html = html_bytes.decode("utf-8", errors="replace")
        text = extract_readable_text_from_html(html)
        if len(text) < 80:
            raise ValueError("Fetched page did not contain enough readable JD text.")
        return text
    except Exception as e:
        if ashby_error is not None:
            raise ValueError(
                "Failed to ingest Ashby JD via both API and HTML fallback. "
                f"API error: {ashby_error}. HTML error: {e}"
            ) from e
        raise


def prompt_for_jd_text() -> str:
    """Prompt the user to paste JD text when URL fetch fails."""
    print("\n" + "=" * 60)
    print("=== JD TEXT INPUT REQUIRED ===")
    print("=" * 60)
    print("Paste the job description text. Submit with an empty line.")
    print("-" * 40)
    lines: list[str] = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line == "":
            break
        lines.append(line)
    return "\n".join(lines).strip()


def infer_role_name(jd_path: Path | None, jd_url: str | None) -> str:
    """Infer role name from JD filename, URL path, or fallback label."""
    if jd_path is not None:
        return core_extract_role_from_jd_filename(jd_path)
    if jd_url:
        parsed = urlparse(jd_url)
        candidate = Path(parsed.path).stem.replace("-", " ").replace("_", " ").strip()
        if candidate:
            return re.sub(r"\s+", " ", candidate)
    return "Tailored Resume"


def get_api_key(cli_key: str | None) -> str:
    """Get API key from CLI arg or environment variable."""
    if cli_key:
        return cli_key
    env_key = os.environ.get("GEMINI_API_KEY")
    if env_key:
        return env_key
    print("ERROR: No API key provided.", file=sys.stderr)
    print("Set GEMINI_API_KEY environment variable or use --api-key argument.", file=sys.stderr)
    sys.exit(1)


def validate_pdf_path(path: Path, label: str) -> Path:
    """Validate that the path exists, is a PDF, and is non-empty."""
    if not path.exists():
        print(f"ERROR: {label} file not found: {path}", file=sys.stderr)
        sys.exit(1)
    if not path.is_file():
        print(f"ERROR: {label} is not a file: {path}", file=sys.stderr)
        sys.exit(1)
    if path.suffix.lower() != ".pdf":
        print(f"ERROR: {label} must be a PDF file (got: {path.suffix})", file=sys.stderr)
        sys.exit(1)
    if path.stat().st_size == 0:
        print(f"ERROR: {label} file is empty: {path}", file=sys.stderr)
        sys.exit(1)
    return path


def extract_text_from_pdf(path: Path) -> str | None:
    """Extract text content from a PDF file."""
    try:
        reader = pypdf.PdfReader(path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)
    except Exception as e:
        print(f"WARNING: Failed to extract text from {path.name}: {e}\nFalling back to PDF input.", file=sys.stderr)
        return None

def load_pdf_content(path: Path) -> str | types.Part:
    """Load PDF content: try text extraction first, fall back to Part."""
    text_content = extract_text_from_pdf(path)
    if text_content and len(text_content.strip()) > 50:
        return f"--- DOCUMENT: {path.name} ---\n{text_content}\n--- END DOCUMENT ---\n"
    
    # Fallback to binary Part if extraction fails or is empty
    with open(path, "rb") as f:
        pdf_bytes = f.read()
    return types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")


def check_for_markdown(text: str) -> list[str]:
    """Wrapper for core markdown detection."""
    return core_check_for_markdown(text)


def send_with_retry(chat: object, message: object, label: str) -> object:
    """Send a message with retry/backoff on quota or rate-limit errors."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            return chat.send_message(message)
        except Exception as e:
            error_msg = str(e).lower()
            is_rate_limit = any(token in error_msg for token in ["quota", "rate", "429", "resource_exhausted"])
            if not is_rate_limit or attempt == MAX_RETRIES:
                raise
            sleep_seconds = min(RETRY_BASE_SECONDS ** attempt, RETRY_MAX_SECONDS)
            print(
                f"WARNING: {label} hit a rate/quota limit. Retrying in {sleep_seconds}s..."
            )
            time.sleep(sleep_seconds)


def print_section(title: str, content: str) -> None:
    """Print a clearly labeled section to terminal."""
    print(f"\n{'=' * 60}")
    print(f"=== {title} ===")
    print("=" * 60)
    print(content)
    print()


def run_architect_draft(
    client: genai.Client,
    jd_content: str | types.Part,
    resume_content: str | types.Part,
    profile_facts_text: str,
) -> tuple[object, str]:
    """Run the Architect agent to create initial draft. Returns (chat, draft_text)."""
    
    draft_prompt = f"""<context>

You will receive two inputs:

1. A document containing multiple tailored versions of my resume-each optimized for roles in data engineering, DevOps, AI/ML, and general software engineering.

2. A target job description I am applying to.



Each resume includes only real, previously written content about my academic, project, and internship experience. No fictional elements are allowed.

Stable profile facts are provided below and should be treated as authoritative:

{profile_facts_text}



</context>



<task>

Your task is to:

- Parse and analyze the job description for keywords, required qualifications, and core responsibilities.

- Cross-reference this with the content in my resume versions.

- Extract the most relevant and high-impact items (work, projects, skills) that align with the job.

- Rewrite and reformat them into a **single one-page resume** tailored to the target role.



You may:

- Rephrase or restructure bullet points for clarity and relevance

- Merge elements from different resume versions

- Adjust section ordering for stronger alignment



You may NOT:

- Fabricate or infer experience

- Add skills, tools, or projects that were not explicitly mentioned in the original documents



The final resume must include:

- Name / Phone / Email / GitHub

- Work Experience

- Project Experience

- Education

- Professional Skills



Make it ready to pass both ATS filters and human recruiters.



</task>



<final_instruction>

Think step-by-step before answering. Prioritize precision, factual relevance, and clear structure.

Output in PLAIN TEXT format only - no markdown syntax.

</final_instruction>"""

    chat = client.chats.create(
        model=MODEL_NAME,
        config=types.GenerateContentConfig(
            system_instruction=ARCHITECT_SYSTEM_INSTRUCTION,
            temperature=0.3,
            top_p=0.95,
            max_output_tokens=8192,
        ),
    )
    
    response = send_with_retry(chat, [jd_content, resume_content, draft_prompt], "Architect draft")
    return chat, core_enforce_plain_text(response.text)


def run_critic(client: genai.Client, draft_resume: str, profile_facts_text: str) -> str:
    """Run the Critic agent to brutally evaluate the draft resume."""
    critic_prompt = build_critic_prompt(draft_resume, profile_facts_text)

    chat = client.chats.create(
        model=MODEL_NAME,
        config=types.GenerateContentConfig(
            temperature=0.4,
            top_p=0.95,
            max_output_tokens=4096,
        ),
    )
    
    response = send_with_retry(chat, critic_prompt, "Critic review")
    return core_enforce_plain_text(response.text)


def run_architect_final(
    architect_chat: object,
    critique: str,
    human_input: str,
    profile_facts_text: str,
) -> str:
    """Run the Architect agent to produce final resume incorporating feedback."""
    
    additional_info_section = ""
    if human_input.strip():
        additional_info_section = f"""
Additional information (if any):

{human_input}"""
    else:
        additional_info_section = "\nAdditional information (if any):\n\n[None provided]"
    
    final_prompt = f"""CRITIQUE FROM RECRUITER:

{critique}

---

Based on all feedback and critiques provided above, rewrite this resume.

Objectives:

- Fully ATS-compliant

- Incorporate role-relevant keywords from the JD

- Clearly reflect required hard skills and soft skills

- Quantify outcomes wherever possible

Rules:

- Do NOT add unverified experiences

- Do NOT soften issues

- Avoid vague or generic phrasing
 
PROFILE FACTS (authoritative):
{profile_facts_text}
{additional_info_section}

Output ONLY the final resume in PLAIN TEXT format - no markdown syntax, no explanations, no preamble."""

    response = send_with_retry(architect_chat, final_prompt, "Architect final")
    return core_enforce_plain_text(response.text)


def get_human_input() -> str:
    """Prompt for and collect human clarifications."""
    print("\n" + "=" * 60)
    print("=== USER INPUT REQUIRED ===")
    print("=" * 60)
    print("Review the draft and critique above.")
    print("Enter any clarifications or additional instructions for the final resume.")
    print("(Press Enter twice to submit, or just Enter to skip)")
    print("-" * 40)
    
    lines: list[str] = []
    empty_count = 0
    
    try:
        while True:
            line = input()
            if line == "":
                if not lines:
                    break
                empty_count += 1
                if empty_count >= 2:
                    break
                lines.append("")
            else:
                empty_count = 0
                lines.append(line)
    except EOFError:
        pass
    
    return "\n".join(lines).strip()


def save_output(filename: str, content: str) -> None:
    """Save content to a file in the workspace directory."""
    path = WORKSPACE_DIR / filename
    path.write_text(content, encoding="utf-8")
    print(f"Saved: {path}")


def extract_role_from_jd_filename(jd_path: Path) -> str:
    """Wrapper for core role extraction."""
    return core_extract_role_from_jd_filename(jd_path)


def save_resume_as_word(content: str, role_name: str) -> Path:
    """Convert plain text resume to Word document and save to tailored_resumes folder."""
    # Ensure folder exists
    TAILORED_RESUMES_FOLDER.mkdir(exist_ok=True)
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Parse and add content
    lines = content.split('\n')
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            # Empty line - add paragraph break
            doc.add_paragraph()
        elif stripped.isupper() and len(stripped) > 2:
            # Section header (all caps)
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
        elif stripped.startswith('- '):
            # Bullet point
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif '|' in stripped and idx < 5:
            # Contact info line (name/email/phone)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            if idx == 0 or (len(stripped) > 10 and stripped[0].isupper()):
                run.bold = True
                run.font.size = Pt(14)
        else:
            # Regular text
            doc.add_paragraph(stripped)
    
    # Generate filename with role and date
    date_str = datetime.datetime.now().strftime("%Y%m%d")
    safe_role = re.sub(r'[^\w\s-]', '', role_name).strip().replace(' ', '_')
    filename = f"{safe_role}_{date_str}.docx"
    filepath = TAILORED_RESUMES_FOLDER / filename
    
    # Save document
    doc.save(filepath)
    print(f"Saved Word document: {filepath}")
    return filepath


def get_application_questions() -> list[str]:
    """Prompt user to enter application questions. Returns list of questions."""
    print("\n" + "=" * 60)
    print("=== APPLICATION QUESTIONS ===")
    print("=" * 60)
    print("Enter your application questions (one per line).")
    print("Press Enter on an empty line when done, or type 'skip' to skip.")
    print("-" * 40)
    
    questions = []
    question_num = 1
    while True:
        try:
            line = input(f"Q{question_num}: ").strip()
        except EOFError:
            break
        
        if line.lower() == 'skip':
            return []
        if not line:
            if questions:  # Only break if we have at least one question
                break
            continue  # Ignore empty first line
        
        questions.append(line)
        question_num += 1
    
    return questions


def extract_length_constraint(question: str) -> str:
    """Wrapper for core length-constraint extraction."""
    return core_extract_length_constraint(question)


def run_application_response(
    client: genai.Client,
    jd_content: str,
    resume: str,
    questions: list[str],
) -> str:
    """Generate high-quality responses to application questions with dynamic length constraints."""
    
    # Build questions with individual constraints
    questions_with_constraints = []
    for i, q in enumerate(questions):
        constraint = extract_length_constraint(q)
        questions_with_constraints.append(
            f"[Question {i+1}] {q}\n   >> LENGTH CONSTRAINT: {constraint}"
        )
    
    questions_formatted = "\n\n".join(questions_with_constraints)
    
    prompt = f"""Based on the latest resume and JD, draft high quality responses to these job application questions/cover letter:

{questions_formatted}

You must strictly follow the below instructions:
- Tone: passionate, sincere, natural, fluent, persuasive, confident
- Length: STRICTLY FOLLOW THE LENGTH CONSTRAINT specified for each question above. This is critical.
- Output format: Plain text only (NO markdown, NO asterisks, NO bullet points)
- Structure each response with a clear header like "QUESTION 1: [question text]" followed by the response

IMPORTANT: If a question specifies "two sentences", your response MUST be exactly two sentences. If it says "100 words", aim for approximately 100 words. Violating length constraints will disqualify the application.

---

JOB DESCRIPTION:
{jd_content}

---

RESUME:
{resume}
"""
    
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.7,
            max_output_tokens=4096,
        ),
    )
    
    return core_enforce_plain_text(response.text)


def save_application_responses(content: str, role_name: str) -> Path:
    """Save application question responses to a txt file."""
    # Ensure folder exists
    APPLICATION_RESPONSES_FOLDER.mkdir(exist_ok=True)
    
    # Generate filename with role and date
    date_str = datetime.datetime.now().strftime("%Y%m%d")
    safe_role = re.sub(r'[^\w\s-]', '', role_name).strip().replace(' ', '_')
    filename = f"{safe_role}_responses_{date_str}.txt"
    filepath = APPLICATION_RESPONSES_FOLDER / filename
    
    # Add header for easy copy-paste
    header = f"""{'=' * 60}
APPLICATION RESPONSES
Role: {role_name}
Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}
{'=' * 60}

"""
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(header + content)
    
    print(f"Saved responses: {filepath}")
    return filepath


def main() -> None:
    """Main entry point."""
    total_start = time.perf_counter()
    
    args = parse_args()
    
    # Get and validate API key
    api_key = get_api_key(args.api_key)
    
    # Resolve resume path
    if args.resume_pdf:
        resume_path = Path(args.resume_pdf)
    else:
        resume_path = DEFAULT_RESUME
        print(f"Using default resume: {resume_path.name}")

    # Resolve profile facts
    profile_facts_path: Path | None = None
    if args.profile_facts:
        profile_facts_path = Path(args.profile_facts)
        if not profile_facts_path.exists():
            print(f"ERROR: --profile-facts file not found: {profile_facts_path}", file=sys.stderr)
            sys.exit(1)
    elif DEFAULT_PROFILE_FACTS.exists():
        profile_facts_path = DEFAULT_PROFILE_FACTS

    try:
        profile_facts = load_profile_facts(profile_facts_path)
    except Exception as e:
        print(f"ERROR: Failed to load profile facts: {e}", file=sys.stderr)
        sys.exit(1)
    profile_facts_text = format_profile_facts_for_prompt(profile_facts)

    jd_path: Path | None = None
    jd_content: str | types.Part
    jd_label: str

    if args.jd_text and args.jd_text.strip():
        jd_label = "inline --jd-text"
        jd_content = f"--- JOB DESCRIPTION (inline text) ---\n{args.jd_text.strip()}\n--- END JOB DESCRIPTION ---\n"
    elif args.jd_url:
        jd_label = args.jd_url
        try:
            with timed_section("JD URL fetch"):
                fetched_jd_text = fetch_jd_text_from_url(args.jd_url)
            jd_content = f"--- JOB DESCRIPTION ({args.jd_url}) ---\n{fetched_jd_text}\n--- END JOB DESCRIPTION ---\n"
        except Exception as e:
            print(f"WARNING: JD URL fetch failed: {e}", file=sys.stderr)
            pasted = prompt_for_jd_text()
            if not pasted:
                print("ERROR: No pasted JD text provided after URL fetch failure.", file=sys.stderr)
                sys.exit(1)
            jd_label = f"{args.jd_url} (fallback pasted text)"
            jd_content = f"--- JOB DESCRIPTION (pasted text) ---\n{pasted}\n--- END JOB DESCRIPTION ---\n"
    else:
        if args.jd_pdf:
            jd_path = Path(args.jd_pdf)
        else:
            jd_path = get_latest_jd_pdf()
            if jd_path is None:
                print(f"ERROR: No PDF files found in JD folder: {JD_FOLDER}", file=sys.stderr)
                print("Provide --jd-url/--jd-text or add a JD PDF.", file=sys.stderr)
                sys.exit(1)
            print(f"Auto-detected latest JD: {jd_path.name}")
        jd_path = validate_pdf_path(jd_path, "Job Description")
        jd_label = str(jd_path)

    # Validate resume path
    resume_path = validate_pdf_path(resume_path, "Resume")
    
    print("-" * 60)
    print(f"Job Description: {jd_label}")
    print(f"Resume: {resume_path}")
    print(f"Profile Facts: {profile_facts_path if profile_facts_path else '[none]'}")
    print(f"Model: {MODEL_NAME}")
    print("-" * 60)
    
    # Load inputs
    if jd_path is not None:
        print("Loading PDF files (parallel extraction)...")
        with timed_section("PDF extraction"):
            jd_content, resume_content = parallel_extract_pdfs(jd_path, resume_path)
    else:
        print("Loading resume PDF...")
        with timed_section("Resume extraction"):
            resume_content = load_pdf_content(resume_path)
    
    # Initialize Gemini client
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        print(f"ERROR: Failed to initialize Gemini client: {e}", file=sys.stderr)
        sys.exit(1)
    
    # === PHASE 1: Architect Draft ===
    print("\n[Phase 1/4] Architect creating initial draft...")
    with Spinner("Gemini drafting resume"), timed_section("Architect Draft API"):
        try:
            architect_chat, draft_response = run_architect_draft(
                client,
                jd_content,
                resume_content,
                profile_facts_text,
            )
        except Exception as e:
            error_msg = str(e).lower()
            if "api key" in error_msg or "authentication" in error_msg or "401" in error_msg:
                print(f"ERROR: Authentication failed. Check your API key.", file=sys.stderr)
            elif "quota" in error_msg or "429" in error_msg:
                print(f"ERROR: API quota exceeded. Try again later.", file=sys.stderr)
            elif "model" in error_msg or "404" in error_msg:
                print(f"ERROR: Model '{MODEL_NAME}' not found or unavailable.", file=sys.stderr)
            else:
                print(f"ERROR: API call failed: {e}", file=sys.stderr)
            sys.exit(1)
    
    print_section("DRAFT RESUME", draft_response)
    
    # Check for markdown in draft
    md_warnings = check_for_markdown(draft_response)
    if md_warnings:
        print("WARNING: Markdown-like formatting detected in draft:")
        for w in md_warnings[:5]:
            print(f"  - {w}")
        print()
    
    # Save draft
    save_output("Draft_Resume.txt", draft_response)
    
    # === PHASE 2: Critic Review ===
    print("\n[Phase 2/4] Critic performing brutal review...")
    with Spinner("Critic analyzing draft"), timed_section("Critic Review API"):
        try:
            critique = run_critic(client, draft_response, profile_facts_text)
        except Exception as e:
            print(f"ERROR: Critic API call failed: {e}", file=sys.stderr)
            sys.exit(1)
    
    print_section("CRITIQUE", critique)
    
    # Save critique
    save_output("Critique.txt", critique)

    parsed_critique = parse_structured_critic(critique)
    save_output("Critique_Structured.json", json.dumps(parsed_critique, indent=2))
    
    # === PHASE 3: Human Input ===
    if has_blocking_critic_issues(parsed_critique):
        print("\nCritic reported blocking user-input issues (Section A).")
        human_input = get_human_input()
        if human_input:
            print(f"\nReceived clarifications ({len(human_input)} chars)")
        else:
            print("\nNo clarifications provided; proceeding with available facts.")
    else:
        human_input = ""
        print("\nNo blocking user-input issues in structured critique; skipping manual pause.")
    
    # === PHASE 4: Final Resume ===
    print("\n[Phase 4/4] Architect producing final resume...")
    with Spinner("Gemini finalizing resume"), timed_section("Architect Final API"):
        try:
            final_resume = run_architect_final(
                architect_chat,
                critique,
                human_input,
                profile_facts_text,
            )
        except Exception as e:
            print(f"ERROR: Final resume API call failed: {e}", file=sys.stderr)
            sys.exit(1)
    
    print_section("FINAL RESUME", final_resume)
    
    # Check for markdown in final
    md_warnings = check_for_markdown(final_resume)
    if md_warnings:
        print("WARNING: Markdown-like formatting detected in final resume:")
        for w in md_warnings[:5]:
            print(f"  - {w}")
        print()
    
    # Save final resume (txt for reference)
    save_output("Final_Resume.txt", final_resume)
    
    # === PHASE 5: Convert to Word and save to tailored_resumes folder ===
    print("\n[Phase 5/6] Converting to Word format...")
    with timed_section("Word conversion"):
        role_name = infer_role_name(jd_path, args.jd_url)
        word_path = save_resume_as_word(final_resume, role_name)
    
    # === PHASE 6: Application Question Responses ===
    print("\n[Phase 6/6] Application Question Responses (Optional)")
    questions = get_application_questions()
    
    responses_path = None
    if questions:
        print(f"\nGenerating responses for {len(questions)} question(s)...")
        with Spinner("Crafting responses"), timed_section("Application Responses API"):
            try:
                # Use extracted text where available; fallback only when binary JD part was used.
                if isinstance(jd_content, str):
                    jd_text = jd_content
                elif jd_path is not None:
                    jd_text = extract_text_from_pdf(jd_path) or "[JD content from PDF]"
                else:
                    jd_text = "[JD content unavailable]"
                responses = run_application_response(client, jd_text, final_resume, questions)
            except Exception as e:
                print(f"ERROR: Failed to generate responses: {e}", file=sys.stderr)
                responses = None
        
        if responses:
            print_section("APPLICATION RESPONSES", responses)
            responses_path = save_application_responses(responses, role_name)
    else:
        print("Skipped application questions.")
    
    # Total time
    total_elapsed = time.perf_counter() - total_start
    
    print("\n" + "=" * 60)
    print(f"COMPLETE! Your tailored resume is saved to:")
    print(f"  - Word: {word_path}")
    print(f"  - Text: {WORKSPACE_DIR / 'Final_Resume.txt'}")
    if responses_path:
        print(f"  - Responses: {responses_path}")
    print(f"  - Total time: {total_elapsed:.2f}s")
    print("=" * 60)


if __name__ == "__main__":
    main()
